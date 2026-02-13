buscador_marcas/
│
├── app.py
├── busqueda_marcas.py
├── requirements.txt
├── todas las marcas uy -para busquedas-09-2025.xlsx
├── Membretada-Uruguay - 2025.docx
import pandas as pd
import unicodedata
from difflib import SequenceMatcher
from docx import Document
import os
import sys

def strip_accents(s):
    if not isinstance(s, str):
        return ""
    return "".join(
        ch for ch in unicodedata.normalize("NFKD", s)
        if not unicodedata.combining(ch)
    )

def norm(s):
    return strip_accents(s).upper().strip()

def sim(a, b):
    return SequenceMatcher(None, a, b).ratio()

archivo = "todas las marcas uy -para busquedas-09-2025.xlsx"

marca = input("👉 Ingresá la marca a buscar: ").strip().upper()
clases = input("👉 Ingresá las clases (ej: 32 o 25 35): ").split()

df = pd.read_excel(archivo)
df.columns = [c.strip().lower() for c in df.columns]

# Filtrado por clases
patron = "|".join([fr"\b{c}\b" for c in clases])
df_filtrado = df[df["clases"].astype(str).str.contains(patron, na=False)].copy()

# Normalización
df_filtrado["denominación_norm"] = df_filtrado["denominación"].apply(norm)
target = norm(marca)

# Similitud con 1 decimal
df_filtrado["sim_target"] = df_filtrado["denominación_norm"].apply(
    lambda s: round(sim(s, target), 1)
)

# Filtro por similitud o coincidencia parcial
mask = (
    df_filtrado["denominación_norm"].str.contains(target[:4], regex=False)
    | (df_filtrado["sim_target"] >= 0.6)
)
result = df_filtrado[mask].copy()

# Clasificación de coincidencia
def tipo_row(row):
    if target in row["denominación_norm"]:
        return "Coincidencia directa"
    elif row["sim_target"] >= 0.8:
        return "Similitud alta (≥80%)"
    else:
        return "Similitud media (60–79%)"

result["tipo_coincidencia"] = result.apply(tipo_row, axis=1)
result.sort_values(by="sim_target", ascending=False, inplace=True)

# Formato de fecha SIN hora
if "fecha" in result.columns:
    result["fecha"] = pd.to_datetime(
        result["fecha"], errors="coerce"
    ).dt.strftime("%d/%m/%Y")

# Generación del Word
doc = Document("Membretada-Uruguay - 2025.docx")
doc.add_heading(
    f"Búsqueda de Disponibilidad – {marca} (Clases {', '.join(clases)})",
    level=1
)

doc.add_paragraph(
    f"Informe de coincidencias directas, parciales y fonéticas para la denominación "
    f"'{marca}' en las clases {', '.join(clases)} del registro de marcas de Uruguay."
)

cols = [
    "número",
    "denominación",
    "clases",
    "titular",
    "status",
    "tipo",
    "fecha",
    "sim_target",
    "tipo_coincidencia",
]
cols = [c for c in cols if c in result.columns]

table = doc.add_table(rows=1, cols=len(cols))
table.style = "Table Grid"

# Encabezados
hdr = table.rows[0].cells
for i, col in enumerate(cols):
    hdr[i].text = col.capitalize()

# Filas
for _, row in result[cols].iterrows():
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = str(val) if pd.notna(val) else ""

# Conclusión
doc.add_page_break()
doc.add_heading("Conclusión", level=1)

altas = (result["sim_target"] >= 0.8).sum()
medias = ((result["sim_target"] >= 0.6) & (result["sim_target"] < 0.8)).sum()
directas = (result["tipo_coincidencia"] == "Coincidencia directa").sum()

if directas > 0:
    riesgo = "ALTO"
elif altas > 0:
    riesgo = "MEDIO"
elif medias > 0:
    riesgo = "MEDIO-BAJO"
else:
    riesgo = "BAJO"

doc.add_paragraph(
    f"• Coincidencias directas: {directas}\n"
    f"• Similitud alta: {altas}\n"
    f"• Similitud media: {medias}\n\n"
    f"Nivel de riesgo: {riesgo}"
)

nombre_salida = f"Disponibilidad_{marca.replace(' ', '_')}_Clases_{'-'.join(clases)}.docx"
doc.save(nombre_salida)

print(f"\n✅ Informe generado: {nombre_salida}")

# Abrir automáticamente el Word
try:
    if sys.platform == "win32":
        os.startfile(nombre_salida)
    elif sys.platform == "darwin":
        os.system(f"open '{nombre_salida}'")
    else:
        os.system(f"xdg-open '{nombre_salida}'")
except Exception as e:
    print(f"No se pudo abrir el archivo automáticamente: {e}")
