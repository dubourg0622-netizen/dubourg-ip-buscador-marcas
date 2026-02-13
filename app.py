import streamlit as st
import pandas as pd
import unicodedata
from difflib import SequenceMatcher
from docx import Document
import os
st.set_page_config(
    page_title="Dubourg IP – Trademark Risk Analysis",
    layout="centered",
)

# -------------------------
# ESTILO MINIMALISTA
# -------------------------

st.markdown("""
    <style>
    .main {
        background-color: #ffffff;
    }
    h1 {
        font-weight: 500;
        letter-spacing: 0.5px;
    }
    .stButton>button {
        background-color: #1C2A39;
        color: white;
        border-radius: 4px;
        height: 40px;
        width: 100%;
        border: none;
    }
    .stButton>button:hover {
        background-color: #2F3E4E;
        color: white;
    }
    .stTextInput>div>div>input {
        border-radius: 4px;
    }
    footer {visibility: hidden;}
    </style>
""", unsafe_allow_html=True)

# ------------------------
# Funciones
# ------------------------

def strip_accents(s):
    if not isinstance(s, str):
        return ""
    return "".join(
        ch for ch in unicodedata.normalize("NFKD", s)
        if not unicodedata.combining(ch)
    )

def norm(s):
    return strip_accents(s).upper().strip()

def sim(a, b):
    return SequenceMatcher(None, a, b).ratio()

# ------------------------
# Interfaz
# ------------------------

st.set_page_config(page_title="Dubourg IP - Buscador de Marcas", layout="centered")

st.title("🔎 Buscador Interno de Marcas")
st.markdown("**Dubourg IP Law Firm – Herramienta interna**")

marca = st.text_input("Denominación a buscar")
clases_input = st.text_input("Clases (ej: 32 o 25 35)")

if st.button("Generar informe"):

    clases = clases_input.split()
    archivo = "todas las marcas uy -para busquedas-09-2025.xlsx"

    df = pd.read_excel(archivo)
    df.columns = [c.strip().lower() for c in df.columns]

    patron = "|".join([fr"\b{c}\b" for c in clases])
    df_filtrado = df[df["clases"].astype(str).str.contains(patron, na=False)].copy()

    df_filtrado["denominación_norm"] = df_filtrado["denominación"].apply(norm)
    target = norm(marca)

    df_filtrado["sim_target"] = df_filtrado["denominación_norm"].apply(
        lambda s: round(sim(s, target), 1)
    )

    mask = (
        df_filtrado["denominación_norm"].str.contains(target[:4], regex=False)
        | (df_filtrado["sim_target"] >= 0.6)
    )

    result = df_filtrado[mask].copy()

    def tipo_row(row):
        if target in row["denominación_norm"]:
            return "Coincidencia directa"
        elif row["sim_target"] >= 0.8:
            return "Similitud alta (≥80%)"
        else:
            return "Similitud media (60–79%)"

    result["tipo_coincidencia"] = result.apply(tipo_row, axis=1)
    result.sort_values(by="sim_target", ascending=False, inplace=True)

    doc = Document("Membretada-Uruguay - 2025.docx")

    doc.add_heading(
        f"Búsqueda de Disponibilidad – {marca} (Clases {', '.join(clases)})",
        level=1
    )

    cols = ["número","denominación","clases","titular","status","fecha","sim_target","tipo_coincidencia"]
    cols = [c for c in cols if c in result.columns]

    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for i, col in enumerate(cols):
        hdr[i].text = col.capitalize()

    for _, row in result[cols].iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    nombre_salida = f"Disponibilidad_{marca}.docx"
    doc.save(nombre_salida)

    with open(nombre_salida, "rb") as file:
        st.download_button(
            label="⬇ Descargar Informe",
            data=file,
            file_name=nombre_salida,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


